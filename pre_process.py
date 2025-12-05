# # pre_processing_backend.py
# from __future__ import annotations

# from datetime import datetime
# import os
# import shutil
# import subprocess
# import tempfile
# import uuid
# import zipfile
# from io import BytesIO
# from pathlib import Path
# from typing import Iterable, Tuple

# import pandas as pd
# import pdfplumber
# from PIL import Image
# from docx import Document
# from docx2pdf import convert
# from fpdf import FPDF


# def save_text_as_pdf(text: str, pdf_file_save_path: str) -> None:
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
#     pdf.multi_cell(0, 6, txt=text)
#     pdf.output(pdf_file_save_path)


# def convert_docx_to_pdf(docx_file_path: str, pdf_file_path: str | None = None) -> None:
#     convert(docx_file_path, pdf_file_path)


# def _ocr_if_needed(file_path: str) -> str:
#     """
#     If the PDF has no text layer, try OCR with ocrmypdf (tesseract required).
#     Returns path to a PDF with text (original or OCR'ed).
#     """
#     contains_text = False
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             if (page.extract_text() or "").strip():
#                 contains_text = True
#                 break

#     if contains_text:
#         return file_path

#     # need OCR
#     if shutil.which("tesseract") is None:
#         raise RuntimeError("Tesseract not found but required for OCR.")
#     if shutil.which("ocrmypdf") is None:
#         raise RuntimeError("OCRmyPDF not found but required for OCR.")

#     ocr_output_path = os.path.join(tempfile.mkdtemp(), f"ocr_{os.path.basename(file_path)}")
#     subprocess.run(
#         ["ocrmypdf", "-l", "deu", "--force-ocr", file_path, ocr_output_path],
#         check=True,
#     )
#     return ocr_output_path


# def _extract_text_from_pdf(pdf_path: str) -> str:
#     text = []
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             t = page.extract_text() or ""
#             text.append(t)
#     return "".join(text)


# def preprocess_files(
#     file_paths: Iterable[str],
#     text_split: int,
#     job_id: str | None = None,
# ) -> Tuple[bytes, pd.DataFrame]:
#     """
#     Pure backend version of 'preprocess_input' + '/download':
#     - Accepts local file paths (pdf/jpg/jpeg/png/txt/docx)
#     - Ensures we have a PDF per input; OCRs image-PDF if needed
#     - Extracts text to 'report'
#     - Builds a CSV with columns: ['report','id','metadata']
#     - Zips the original/converted PDFs (renamed to <id>.<ext>) + CSV 'preprocessed_<job_id>.csv'
#     Returns: (zip_bytes, df_split)
#     """
#     file_paths = list(file_paths)
#     if job_id is None:
#         job_id = datetime.now().strftime("%Y%m%d%H%M") + "-" + uuid.uuid4().hex[:8]

#     merged_rows = []
#     kept_pdf_paths = []   # these are the file paths we will zip alongside
#     temp_dirs = []        # for cleanup if needed

#     try:
#         for file_path in file_paths:
#             file_path = str(file_path)
#             if file_path.lower().endswith((".pdf", ".jpg", ".jpeg", ".png")):
#                 # Convert image to PDF if needed
#                 if not file_path.lower().endswith(".pdf"):
#                     td = tempfile.mkdtemp()
#                     temp_dirs.append(td)
#                     pdf_output_path = os.path.join(td, f"pdf_{os.path.basename(file_path)}.pdf")
#                     image = Image.open(file_path).convert("RGB")
#                     image.save(pdf_output_path)
#                     file_path = pdf_output_path

#                 # OCR if PDF has no text layer
#                 pdf_with_text = _ocr_if_needed(file_path)
#                 text = _extract_text_from_pdf(pdf_with_text)

#                 merged_rows.append({"report": text, "filepath": pdf_with_text})
#                 kept_pdf_paths.append(pdf_with_text)

#             elif file_path.lower().endswith(".txt"):
#                 with open(file_path, "r", encoding="utf-8") as f:
#                     text = f.read()
#                 td = tempfile.mkdtemp()
#                 temp_dirs.append(td)
#                 pdf_file_save_path = os.path.join(td, f"ocr_{Path(file_path).stem}.pdf")
#                 save_text_as_pdf(text, pdf_file_save_path)
#                 merged_rows.append({"report": text, "filepath": pdf_file_save_path})
#                 kept_pdf_paths.append(pdf_file_save_path)

#             elif file_path.lower().endswith(".docx"):
#                 doc = Document(file_path)
#                 doc_text = "\n".join(p.text for p in doc.paragraphs)

#                 td = tempfile.mkdtemp()
#                 temp_dirs.append(td)
#                 pdf_file_save_path = os.path.join(td, f"ocr_{Path(file_path).stem}.pdf")
#                 convert_docx_to_pdf(file_path, pdf_file_save_path)

#                 merged_rows.append({"report": doc_text, "filepath": pdf_file_save_path})
#                 kept_pdf_paths.append(pdf_file_save_path)

#             else:
#                 raise ValueError(f"Unsupported file format: {file_path}")

#         df = pd.DataFrame(merged_rows)

#         # filename column (remove leading 'ocr_' if present)
#         def remove_ocr_prefix(filename: str) -> str:
#             base = os.path.basename(filename)
#             return base[4:] if base.startswith("ocr_") else base

#         df["filename"] = df["filepath"].apply(lambda p: remove_ocr_prefix(os.path.basename(p)))
#         df["id"] = df.apply(lambda x: f'{x["filename"]}${uuid.uuid4()}', axis=1)

#         # metadata with preprocessing timestamp
#         df["metadata"] = df.apply(
#             lambda _: {"preprocessing": {"date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}},
#             axis=1,
#         )

#         # split report by max length
#         split_rows = []
#         for _, row in df.iterrows():
#             text = row["report"]
#             if len(text) > text_split:
#                 num_splits = (len(text) + text_split - 1) // text_split
#                 for i in range(num_splits):
#                     r = row.copy()
#                     r["report"] = text[i * text_split : (i + 1) * text_split]
#                     r["id"] = f'{row["id"]}_{i}'
#                     split_rows.append(r)
#             else:
#                 split_rows.append(row)

#         df_split = pd.DataFrame(split_rows)

#         # build zip buffer: include PDFs (renamed <id>.<ext>) + CSV
#         zip_buffer = BytesIO()
#         with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zipf:
#             for file_path, file_id in zip(df["filepath"].tolist(), df["id"].tolist()):
#                 ext = os.path.basename(file_path).split(".")[-1]
#                 arcname = f"{file_id}.{ext}"   # e.g., <filename>$<uuid>.pdf
#                 zipf.write(file_path, arcname=arcname)

#             with tempfile.TemporaryDirectory() as tmpd:
#                 csv_filename = f"preprocessed_{job_id}.csv"
#                 csv_path = os.path.join(tmpd, csv_filename)
#                 # 只保留 report/id/metadata 三列写入 CSV（老版对齐）
#                 df_split[["report", "id", "metadata"]].to_csv(csv_path, index=False)
#                 zipf.write(csv_path, arcname=csv_filename)

#         zip_buffer.seek(0)
#         return zip_buffer.getvalue(), df_split


#     finally:
#         # you can keep temporary dirs for debugging by commenting this out
#         for td in temp_dirs:
#             shutil.rmtree(td, ignore_errors=True)


# if __name__ == "__main__":
#     import sys

#     try:
#         file_input = input("Enter file paths (separated by space): ").strip()
#         if not file_input:
#             print("No input files provided.")
#             sys.exit(1)
#         file_paths = file_input.split()

#         text_split_str = input("Enter text_split (max chars per chunk): ").strip()
#         text_split = int(text_split_str) if text_split_str else 3000

#         zip_bytes, df = preprocess_files(file_paths, text_split=text_split)

#         zip_out = "preprocessed_output.zip"
#         csv_out = "preprocessed_output.csv"

#         with open(zip_out, "wb") as f:
#             f.write(zip_bytes)
#         df.to_csv(csv_out, index=False)

#         print(f"[OK] Wrote zip to: {zip_out}")
#         print(f"[OK] Wrote CSV to: {csv_out}")
#         print("[HEAD]")
#         print(df.head().to_string(index=False))

#     except Exception as e:
#         print(f"[ERROR] {e}", file=sys.stderr)
#         sys.exit(1)



# pre_processing_backend.py  — text-only lite version (no PDF outputs)
from __future__ import annotations

from datetime import datetime
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Iterable, Tuple, List, Optional

import pandas as pd
import pdfplumber
from PIL import Image
from docx import Document

# -----------------------------
# Helpers
# -----------------------------

def _extract_paths_from_zip(zip_path: str,
                            allowed_exts=(".pdf", ".txt", ".docx", ".jpg", ".jpeg", ".png")) -> Tuple[str, List[str]]:
    """
    Extract allowed files from a zip into a temp directory and return (temp_dir, extracted_file_paths).
    Caller must cleanup temp_dir.
    """
    if not zip_path.lower().endswith(".zip"):
        raise ValueError(f"Not a zip file: {zip_path}")

    tmpd = tempfile.mkdtemp()
    extracted_files: List[str] = []
    with zipfile.ZipFile(zip_path, "r") as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            if not info.filename.lower().endswith(allowed_exts):
                continue
            out_path = os.path.join(tmpd, info.filename)
            os.makedirs(os.path.dirname(out_path), exist_ok=True)
            with zf.open(info) as src, open(out_path, "wb") as dst:
                dst.write(src.read())
            extracted_files.append(out_path)
    return tmpd, extracted_files


def _pdf_has_text(pdf_path: str) -> bool:
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                if (page.extract_text() or "").strip():
                    return True
        return False
    except Exception:
        return False


def _ocr_pdf_to_text(pdf_path: str, lang: str = "eng") -> str:
    """
    If ocrmypdf & tesseract are available, OCR to a temp PDF then extract text, else return empty string.
    We do not persist the OCR-ed PDF; only return extracted text.
    """
    if shutil.which("tesseract") is None or shutil.which("ocrmypdf") is None:
        return ""  # silently skip OCR if not available
    td = tempfile.mkdtemp()
    try:
        ocr_pdf = os.path.join(td, f"ocr_{os.path.basename(pdf_path)}")
        subprocess.run(["ocrmypdf", "-l", lang, "--force-ocr", pdf_path, ocr_pdf], check=True)
        # extract text
        text_parts: List[str] = []
        with pdfplumber.open(ocr_pdf) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                text_parts.append(t)
        return "".join(text_parts)
    except Exception:
        return ""
    finally:
        shutil.rmtree(td, ignore_errors=True)


# def _image_to_text(img_path: str, lang: str = "eng") -> str:
#     """
#     OCR image to text.
#     Prefer pytesseract if available, otherwise convert to one-page PDF and call _ocr_pdf_to_text.
#     """
#     try:
#         import pytesseract  # optional dependency
#         img = Image.open(img_path).convert("RGB")
#         return pytesseract.image_to_string(img, lang=lang)
#     except Exception:
#         # fallback: save image as temp PDF then run OCR PDF pipeline
#         td = tempfile.mkdtemp()
#         try:
#             pdf_path = os.path.join(td, f"tmp_{Path(img_path).stem}.pdf")
#             Image.open(img_path).convert("RGB").save(pdf_path)
#             return _ocr_pdf_to_text(pdf_path, lang=lang)
#         finally:
#             shutil.rmtree(td, ignore_errors=True)


def _pdf_to_text(pdf_path: str, ocr_lang: str = "eng") -> str:
    """
    Extract text from a PDF; if it has no text layer, try OCR.
    """
    if _pdf_has_text(pdf_path):
        parts: List[str] = []
        with pdfplumber.open(pdf_path) as pdf:
            for p in pdf.pages:
                parts.append(p.extract_text() or "")
        return "".join(parts)
    # try OCR
    return _ocr_pdf_to_text(pdf_path, lang=ocr_lang)


def _docx_to_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)


def _read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# -----------------------------
# Core APIs (no PDF outputs)
# -----------------------------

def preprocess_files(
    file_paths: Iterable[str],
    text_split: int,
    job_id: Optional[str] = None,
    ocr_lang: str = "eng",
) -> Tuple[bytes, pd.DataFrame]:
    """
    Text-only preprocessor:
    - Accept paths of .pdf/.txt/.docx/.jpg/.jpeg/.png
    - Extract text to a DataFrame with columns: ['report','id','metadata','source']
    - Split 'report' by char length
    - Return (zip_bytes_only_containing_csv, df_split)
    ZIP now only contains a CSV named 'preprocessed_<job_id>.csv'
    """
    file_paths = list(file_paths)
    if job_id is None:
        job_id = datetime.now().strftime("%Y%m%d%H%M") + "-" + uuid.uuid4().hex[:8]

    rows: List[dict] = []
    for path in file_paths:
        path = str(path)
        ext = Path(path).suffix.lower()
        try:
            if ext == ".txt":
                text = _read_text_file(path)
            elif ext == ".pdf":
                text = _pdf_to_text(path, ocr_lang)
            elif ext == ".docx":
                text = _docx_to_text(path)
            # elif ext in (".jpg", ".jpeg", ".png"):
            #     text = _image_to_text(path, lang=ocr_lang)
            else:
                raise ValueError(f"Unsupported file format: {path}")

            rows.append({"report": text, "source": os.path.basename(path)})
        except Exception as e:
            # 如果单个文件失败，写一条空文本但保留来源，避免全批中断（也便于后续排查）
            rows.append({"report": "", "source": os.path.basename(path)})

    df = pd.DataFrame(rows)

    # add id & metadata
    df["id"] = df.apply(lambda x: f'{x["source"]}${uuid.uuid4()}', axis=1)
    df["metadata"] = df.apply(
        lambda _: {"preprocessing": {"date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}},
        axis=1,
    )

    # split long reports
    split_rows: List[dict] = []
    for _, row in df.iterrows():
        text = row["report"] or ""
        if len(text) > text_split:
            num = (len(text) + text_split - 1) // text_split
            for i in range(num):
                r = row.copy()
                r["report"] = text[i * text_split : (i + 1) * text_split]
                r["id"] = f'{row["id"]}_{i}'
                split_rows.append(r)
        else:
            split_rows.append(row)

    df_split = pd.DataFrame(split_rows)

    # build zip buffer: ONLY CSV
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zipf:
        with tempfile.TemporaryDirectory() as tmpd:
            csv_filename = f"preprocessed_{job_id}.csv"
            csv_path = os.path.join(tmpd, csv_filename)
            df_split[["report", "id", "metadata"]].to_csv(csv_path, index=False)
            zipf.write(csv_path, arcname=csv_filename)

    zip_buffer.seek(0)
    return zip_buffer.getvalue(), df_split


def preprocess_zip(zip_path: str, text_split: int, job_id: Optional[str] = None, ocr_lang: str = "eng") -> Tuple[bytes, pd.DataFrame]:
    tmpd, paths = _extract_paths_from_zip(zip_path)
    try:
        return preprocess_files(paths, text_split=text_split, job_id=job_id, ocr_lang=ocr_lang)
    finally:
        shutil.rmtree(tmpd, ignore_errors=True)


if __name__ == "__main__":
    import sys
    try:
        file_input = input("Enter a zip path OR file paths (space-separated): ").strip()
        if not file_input:
            print("No input provided.")
            sys.exit(1)
        inputs = file_input.split()

        text_split_str = input("Enter text_split (max chars per chunk): ").strip()
        text_split = int(text_split_str) if text_split_str else 3000

        ocr_lang = input("OCR language (default 'eng'): ").strip() or "eng"

        if len(inputs) == 1 and inputs[0].lower().endswith(".zip"):
            zip_bytes, df = preprocess_zip(inputs[0], text_split=text_split, ocr_lang=ocr_lang)
        else:
            zip_bytes, df = preprocess_files(inputs, text_split=text_split, ocr_lang=ocr_lang)

        # outputs
        zip_out = "preprocessed_output.zip"
        csv_out = "preprocessed_output.csv"
        with open(zip_out, "wb") as f:
            f.write(zip_bytes)
        df.to_csv(csv_out, index=False)

        print(f"[OK] Wrote ZIP (CSV only): {zip_out}")
        print(f"[OK] Wrote CSV preview:   {csv_out}")
        print("[HEAD]")
        print(df.head().to_string(index=False))
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)

